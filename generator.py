import os
import uuid
from docx import Document
from docx.shared import Cm
import subprocess

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, 'shablon', 'Abiturient_qayd_varaqasi.docx')
OUTPUT_DIR = os.path.join(BASE_DIR, 'generated_docs')

# Fayllar saqlanadigan papkani yaratish
os.makedirs(OUTPUT_DIR, exist_ok=True)


def insert_image(doc, image_path):
    """
    DOCX hujjatidagi {{image}} joyini qidiradi va rasm joylaydi.
    """
    # Paragraph ichidagi run'lar
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if '{{image}}' in run.text:
                run.clear()
                run.add_picture(image_path, width=Cm(3), height=Cm(4))
                return True

    # Table ichidagi run'lar
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if '{{image}}' in run.text:
                            run.clear()
                            run.add_picture(image_path, width=Cm(3), height=Cm(4))
                            return True

    # Header / Footer ichida
    for section in doc.sections:
        for area in [section.header, section.footer]:
            for paragraph in area.paragraphs:
                for run in paragraph.runs:
                    if '{{image}}' in run.text:
                        run.clear()
                        run.add_picture(image_path, width=Cm(3), height=Cm(4))
                        return True

    return False


def generate_docx_and_pdf(user_data):
    """
    DOCX shablonni to‘ldiradi, rasm qo‘shadi, va PDF ga o‘tkazadi.
    """
    doc = Document(TEMPLATE_PATH)

    # Matnni almashtirish
    replace_dict = {
        '{{id}}': str(user_data.get('id', '')),
        '{{fio}}': user_data.get('fio', ''),
        '{{passport}}': user_data.get('pasport', ''),
        '{{jshshir}}': user_data.get('jshshir', ''),
        '{{tugulgan_sana}}': user_data.get('tugulgan_sana', ''),
        '{{jinsi}}': user_data.get('jinsi', ''),
        '{{manzili}}': user_data.get('manzili', ''),
        '{{talim_muassasasi}}': user_data.get('talim_muassasasi', ''),
    }

    for paragraph in doc.paragraphs:
        for key, value in replace_dict.items():
            if key in paragraph.text:
                inline = paragraph.runs
                for run in inline:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))

    # Jadval ichidagilarni ham almashtirish
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replace_dict.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, str(value))

    # Rasm joylash
    image_path = user_data.get('rasm_path')
    if not image_path or not os.path.exists(image_path):
        print("❌ Rasm mavjud emas:", image_path)
        return None

    if not insert_image(doc, image_path):
        print("❌ Rasm joylashda xatolik: {{image}} joyi topilmadi.")
        return None

    # Fayl nomlari
    unique_name = f"{user_data['pasport']}_{uuid.uuid4().hex[:6]}"
    docx_filename = os.path.join(OUTPUT_DIR, f"{unique_name}.docx")
    pdf_filename = os.path.join(OUTPUT_DIR, f"{unique_name}.pdf")

    # Saqlash
    doc.save(docx_filename)

    # PDF ga o‘tkazish
    try:
        subprocess.run([
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            docx_filename,
            '--outdir', OUTPUT_DIR
        ], check=True, env={**os.environ, 'HOME': '/tmp', 'LD_LIBRARY_PATH': ''})
    except subprocess.CalledProcessError as e:
        print("❌ LibreOffice konvertatsiya xatoligi:", e)
        return None

    # DOCX faylni o‘chirish
    if os.path.exists(docx_filename):
        os.remove(docx_filename)

    return pdf_filename  # Bot orqali yuboriladi
